#!/usr/bin/env python3.11
"""Export all summary Markdown files to a single DOCX for Google Docs import."""
import argparse
import glob
import json
import os
import re
import sqlite3
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SUMMARY_DIR = os.path.join(SCRIPT_DIR, "college_summary")
DB_PATH = os.path.join(SCRIPT_DIR, "college_gopher.db")


def is_email_ready(cid, conn):
    """Check if school has coaches with email, season record, and roster >= 3."""
    coaches_row = conn.execute("SELECT data_json FROM digest WHERE cid=? AND digest_type='coaches' AND status='ok'", (cid,)).fetchone()
    coaches = json.loads(coaches_row[0]) if coaches_row and coaches_row[0] else []
    if not any(c.get('email') for c in coaches if c):
        return False
    season_row = conn.execute("SELECT data_json FROM digest WHERE cid=? AND digest_type='season' AND status='ok'", (cid,)).fetchone()
    season = json.loads(season_row[0]) if season_row and season_row[0] else None
    if not season:
        gc = conn.execute("SELECT data_json FROM grounded_cache WHERE cid=? AND cache_type='season'", (cid,)).fetchone()
        season = json.loads(gc[0]) if gc else None
    if not season or not season.get('record'):
        return False
    roster_row = conn.execute("SELECT data_json FROM digest WHERE cid=? AND digest_type='roster' AND status='ok'", (cid,)).fetchone()
    roster = json.loads(roster_row[0]) if roster_row and roster_row[0] else []
    return len(roster) >= 3


def _add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tag = OxmlElement('w:bookmarkStart')
    tag.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    tag.set(qn('w:name'), bookmark_name)
    paragraph._p.append(tag)
    tag2 = OxmlElement('w:bookmarkEnd')
    tag2.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    paragraph._p.append(tag2)


def _add_internal_link(paragraph, bookmark_name, text):
    """Add an internal hyperlink (link to bookmark) in a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs_with_links(paragraph, text):
    """Add text to paragraph, converting [text](url) to hyperlinks and **bold**."""
    parts = re.split(r'(\[.+?\]\(.+?\)|\*\*.+?\*\*)', text)
    for part in parts:
        link_match = re.match(r'\[(.+?)\]\((.+?)\)', part)
        bold_match = re.match(r'\*\*(.+?)\*\*', part)
        if link_match:
            add_hyperlink(paragraph, link_match.group(2), link_match.group(1))
        elif bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
        elif part:
            paragraph.add_run(part)


def add_md_to_doc(doc, md):
    """Convert Markdown to DOCX paragraphs."""
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        # Table
        elif line.startswith("|") and "|---" not in line:
            # Collect table rows
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if "|---" not in lines[i]:
                    cells = [c.strip() for c in lines[i].split("|")[1:-1]]
                    rows.append(cells)
                i += 1
            i -= 1  # back up one
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        # Strip markdown bold
                        cell = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                        table.rows[r_idx].cells[c_idx].text = cell
        # List items
        elif line.startswith("- "):
            content = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_links(p, content)
        elif line.startswith("  - "):
            content = line[4:]
            p = doc.add_paragraph(style='List Bullet 2')
            add_runs_with_links(p, content)
        # Link line
        elif line.startswith("[") and "](" in line:
            m = re.match(r"\[(.+?)\]\((.+?)\)", line)
            if m:
                p = doc.add_paragraph()
                add_hyperlink(p, m.group(2), m.group(1))
        # Bold line (like **Team ID:** `xxx`)
        elif line.startswith("**"):
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            content = re.sub(r'`(.+?)`', r'\1', content)
            doc.add_paragraph(content)
        # Horizontal rule
        elif line.startswith("---"):
            pass  # skip
        # Regular text
        elif line.strip():
            p = doc.add_paragraph()
            add_runs_with_links(p, line)

        i += 1


def main():
    parser = argparse.ArgumentParser(description="Export summaries to DOCX for Google Docs")
    parser.add_argument("-o", "--output", default=os.path.join(SCRIPT_DIR, "all_colleges_summary.docx"))
    args = parser.parse_args()

    files = sorted(glob.glob(os.path.join(SUMMARY_DIR, "*.md")))
    if not files:
        print("No summary files found in college_summary/")
        sys.exit(1)

    # Filter to email-ready schools only
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    ready_files = []
    for f in files:
        cid = os.path.basename(f).replace(".md", "")
        if is_email_ready(cid, conn):
            score = conn.execute("SELECT score_total FROM school_scores WHERE cid=?", (cid,)).fetchone()
            ready_files.append((score[0] if score else 0, f))

    # Sort by score descending
    ready_files.sort(key=lambda x: -x[0])
    files = [f for _, f in ready_files]
    conn.close()

    if not files:
        print("No email-ready schools found.")
        sys.exit(1)

    doc = Document()
    doc.add_heading("College WBB Program Summaries", level=0)

    # Re-open conn for TOC
    conn2 = sqlite3.connect(DB_PATH)
    conn2.row_factory = sqlite3.Row

    # Build TOC grouped by classification
    groups = {'Likely Interest D1': [], 'Great Fit': [], 'Special School': [], '': []}
    cid_list = []
    for score_val, f in ready_files:
        cid = os.path.basename(f).replace(".md", "")
        cid_list.append(cid)
        r = conn2.execute("SELECT classification, college FROM school_scores WHERE cid=?", (cid,)).fetchone()
        cls = r['classification'] if r else ''
        college = r['college'] if r else cid
        if cls not in groups:
            groups[''].append((score_val, college, cid))
        else:
            groups[cls].append((score_val, college, cid))
    conn2.close()

    # Write TOC with internal links
    doc.add_heading("Table of Contents", level=1)
    for section, label in [('Likely Interest D1', '🎯 Likely Interest — D1 Mid-Majors'),
                           ('Great Fit', '✅ Great Fit — D2/NAIA'),
                           ('Special School', '⭐ Special Schools'),
                           ('Beach School', '🏖️ Beach Schools'),
                           ('', '📋 Other Schools')]:
        schools_in_group = groups.get(section, [])
        if schools_in_group:
            doc.add_heading(label, level=2)
            for score_val, college, cid in schools_in_group:
                p = doc.add_paragraph(style='List Bullet')
                _add_internal_link(p, cid, f"{college} ({score_val:.0f})")

    doc.add_page_break()

    for idx, (_, f) in enumerate(ready_files):
        cid = os.path.basename(f).replace(".md", "")
        with open(f) as fh:
            md = fh.read()
        # Strip Data Provenance section
        md = re.split(r"\n---\n## Data Provenance", md)[0]
        # Add bookmark for TOC link
        bookmark_p = doc.add_paragraph()
        _add_bookmark(bookmark_p, cid)
        add_md_to_doc(doc, md)
        if idx < len(ready_files) - 1:
            doc.add_page_break()

    doc.save(args.output)
    print(f"Wrote {args.output}: {len(ready_files)} schools")


if __name__ == "__main__":
    main()
